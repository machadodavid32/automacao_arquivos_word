from docx import Document

desafio = Document()

desafio.add_heading('Relatorio de vendas', 0)

desafio.add_paragraph(
'Neste mês de agosto foram realizadas um total de 10 vendas de notebooks.Segue em anexo a tabela com os dados de cada venda realizada.'
)

desafio.add_heading('Vendas de Agosto', 1)

registros = [['Modelo 1', 'R$1500.00', '25/08/2022', 'Robert', 'Loja 1'],
             ['Modelo 2', 'R$3500.00', '23/08/2022', 'Downey', 'Loja 2'],
             ['Modelo 3', 'R$1400.00', '25/05/2022', 'Junior', 'Loja 3'],
             ['Modelo 4', 'R$3500.00', '24/08/2022', 'Iron', 'Loja 4'] 
]

tabela = desafio.add_table(rows=1, cols=5)
cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Modelo'
cabecalho[1].text = 'Preço'
cabecalho[2].text = 'Data'
cabecalho[3].text = 'Vendedor'
cabecalho[4].text = 'Loja'

for modelo, preco, data, vendedor, loja in registros:
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = modelo
    linha_atual[1].text = preco
    linha_atual[2].text = data
    linha_atual[3].text = vendedor
    linha_atual[4].text = loja
    
    
desafio.add_paragraph('')  # pra ficar uma quebra de linha    

desafio.add_paragraph('Para as vendas deste mês, o funcionário Robert foi o funcionário com a maior quantidade de vendas diretas.')

desafio.add_paragraph('O lucro total gerado na loja 1 foi de R$37000.00')

desafio.save('desafio.docx')

