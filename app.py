# word
from docx import Document # nos permite criar arquivos word
from docx.shared import Cm  # para importar imagens no tamanho em centimetros correto

documento = Document()

# para add um titulo
documento.add_heading('Titulo do documento', 0) # O 0 significa o nível do texto, são varios niveis, ver na documentação



# obs: Toda vez que salvar com o mesmo nome, o arquivo será sobescrevido. Caso queira manter o arquivo atual, mude o nome do save.


# Adicionando um paragrafo
paragrafo = documento.add_paragraph('Um parágrafo simples')
paragrafo.add_run(' e super importante ').bold = True  # deixando esta frase em negrito.
paragrafo.add_run(' do autor ') # sem nenhum tipo de formatação
paragrafo.add_run(' jonathan ').italic = True  # deixando em italico


# Adicionar headings(cabeçalho)
documento.add_heading('Titulo Nível 1', level=1)
documento.add_heading('Titulo Nível 2', level=2)
documento.add_heading('Titulo Nível 3', level=3)
documento.add_heading('Titulo Nível 4', level=4)

# Formatação de estilo
# Formatação de estilo
documento.add_paragraph('Formatação "No Spacing"',style='No Spacing')
documento.add_paragraph('Formatação "Heading1"', style='Heading 1')
documento.add_paragraph('Formatação "Heading 2"', style='Heading 2')
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3')
documento.add_paragraph('Formatação "Title"', style='Title')
documento.add_paragraph('Formatação "Subtitle"', style='Subtitle')
documento.add_paragraph('Formatação "Quote"', style='Quote')
documento.add_paragraph('Formatação "Intense Quote"', style='Intense Quote')
documento.add_paragraph('Formatação "List Paragraph"', style='List Paragraph')
documento.add_paragraph('Primeiro item em uma lista com pontos', style='List Bullet')
documento.add_paragraph('primeiro item em uma lista numerada', style='List Number')


# Adicionando imagens
documento.add_picture('computador.png', width=Cm(5.25))  # o segundo parâmetro é pra definir o tamanho da imagem

# Adicionando tabelas - modo chato
"""
tabela = documento.add_table(rows=3, cols=2) # criando tabela com 3 linhas e 2 colunas
celula00 = tabela.cell(0,0) # celula, linha zero coluna zero
celula00.text = 'Nome' # tudo acima para criar uma unica celula
celula01 = tabela.cell(0,1) # Linha 0 e coluna 1 (formato de indice, ou seja, 0 é o primeiro)
celula01.text = 'Idade'
# Acima é um processo chato e lento para adicionar conteúdo na tabela criada
"""

# Adicionando tabela, modo correto
registros = [
    [3, '101', 'maça'],
    [7, '422', 'ovos'],
    [4, '631', 'Banana']
]

tabela = documento.add_table(rows=1, cols=3) # uma linha pois o conteúdo será gerado dinamicamente
cabecalho = tabela.rows[0].cells # Criando o nome dos cabeçalhos
cabecalho[0].text='Quantidade'
cabecalho[1].text ='Id'
cabecalho[2].text='Descrição'

for quantidade, id, descricao in registros:  # Aqui estamos populando a tabela
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = str(quantidade)  # passando para string para funcionar no word
    linha_atual[1].text = id
    linha_atual[2].text = descricao
    
    
# Salvando o documento criado
documento.save('demo.docx')  # o nome
